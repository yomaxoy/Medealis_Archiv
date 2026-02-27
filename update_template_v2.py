"""
Robustes Hilfsskript zum Aktualisieren des Wareneingangskontrolle-Templates
Ersetzt [[itemnr]] durch [[artikel]] im Word-Dokument
"""
from docx import Document
from pathlib import Path
import shutil

template_path = Path(r"C:\Users\krueg\OneDrive\Desktop\Medealis Archiv\resources\templates\Fo0113_Wareneingangskontrolle.docx")
backup_path = Path(r"C:\Users\krueg\OneDrive\Desktop\Medealis Archiv\resources\templates\Fo0113_Wareneingangskontrolle.docx.backup2")

print(f"Lade Template: {template_path}")

# Backup erstellen
shutil.copy2(template_path, backup_path)
print(f"Backup erstellt: {backup_path}")

# Template laden
doc = Document(template_path)

# Durchsuche alle Tabellen und Zellen
replacements_made = 0
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            # Kombiniere alle Runs zu einem Text
            for paragraph in cell.paragraphs:
                full_text = ''.join([run.text for run in paragraph.runs])

                if '[[itemnr]]' in full_text:
                    print(f"Tabelle {table_idx}, Zeile {row_idx}, Zelle {cell_idx}")
                    print(f"  Original: '{full_text}'")

                    # Neue Text mit Ersetzung
                    new_text = full_text.replace('[[itemnr]]', '[[artikel]]')
                    print(f"  Neu: '{new_text}'")

                    # Lösche alle bestehenden Runs
                    for run in paragraph.runs:
                        run.text = ''

                    # Füge neuen Text als ersten Run hinzu
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)

                    replacements_made += 1

# Speichern
if replacements_made > 0:
    doc.save(template_path)
    print(f"\nTemplate aktualisiert! ({replacements_made} Ersetzungen)")
    print(f"Gespeichert: {template_path}")
else:
    print("\nKeine [[itemnr]] Platzhalter gefunden!")

print("\nVerifikation:")
doc2 = Document(template_path)
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if '[[artikel]]' in text:
                print(f"  OK: [[artikel]] gefunden: '{text}'")
            if '[[itemnr]]' in text:
                print(f"  FEHLER: [[itemnr]] noch vorhanden: '{text}'")
