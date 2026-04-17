"""
Template Manager - Zentrale Template-Verwaltung

Übernimmt die Template-Loading und Verwaltungs-Logik aus WordTemplateService.
SINGLE SOURCE OF TRUTH für alle Template-Operationen.

PERFORMANCE: Nutzt optimierte TemplateCache aus shared/performance
"""

import logging
import sys
from pathlib import Path
from typing import Dict, Optional, List, Any
from datetime import datetime

# External dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - Word template functionality disabled")

# Internal imports
from .document_types import DocumentType, TemplateInfo, TEMPLATE_REGISTRY, get_template_info
from .generation_models import ValidationResult

# PERFORMANCE: Optimierte TemplateCache verwenden (Fallback auf lokale Implementation)
try:
    from warehouse.shared.performance.document_pipeline import TemplateCache as OptimizedTemplateCacheBase
    USE_OPTIMIZED_CACHE = True
    logger = logging.getLogger(__name__)
    logger.info("Using OPTIMIZED TemplateCache from shared/performance")
except ImportError:
    USE_OPTIMIZED_CACHE = False
    OptimizedTemplateCacheBase = None
    logger = logging.getLogger(__name__)
    logger.warning("Optimized TemplateCache not available - using FALLBACK cache")


# Adapter für OptimizedTemplateCache um alte API zu unterstützen
class OptimizedTemplateCacheAdapter:
    """
    Adapter für OptimizedTemplateCache aus shared/performance.

    Macht die optimierte Cache-Implementation kompatibel mit der
    bestehenden TemplateManager API (get/put statt get_template).
    """

    def __init__(self, max_size: int = 20):
        self._optimized_cache = OptimizedTemplateCacheBase(max_size=max_size)
        self._cache_key_to_path: Dict[str, tuple] = {}  # Maps cache_key -> (name, path)
        logger.debug(f"Initialized OptimizedTemplateCacheAdapter (max_size={max_size})")

    def get(self, cache_key: str) -> Optional[Document]:
        """Holt Template aus optimiertem Cache (alte API)."""
        try:
            if cache_key in self._cache_key_to_path:
                name, path = self._cache_key_to_path[cache_key]
                return self._optimized_cache.get_template(name, path)
            return None
        except Exception as e:
            logger.warning(f"Error getting template from optimized cache: {e}")
            return None

    def put(self, cache_key: str, template: Document, template_info: TemplateInfo):
        """Speichert Template-Pfad für späteren Abruf (alte API)."""
        try:
            # Speichere Mapping für get()
            template_path = self._resolve_template_path(template_info)
            self._cache_key_to_path[cache_key] = (template_info.name, template_path)
            # Template wird beim nächsten get() automatisch gecacht via @ttl_cache
        except Exception as e:
            logger.warning(f"Error putting template to optimized cache: {e}")

    def _resolve_template_path(self, template_info: TemplateInfo) -> Path:
        """Löst Template-Pfad auf."""
        # Muss mit TemplateManager._get_template_path konsistent sein
        return Path(__file__).parent.parent.parent.parent.parent.parent / "resources" / "templates" / template_info.filename

    def clear(self):
        """Leert Cache."""
        self._cache_key_to_path.clear()
        if hasattr(self._optimized_cache, 'clear_old_templates'):
            self._optimized_cache.clear_old_templates(max_age_seconds=0)

    def get_stats(self) -> Dict[str, Any]:
        """Gibt Cache-Statistiken zurück."""
        return {
            'cached_keys': len(self._cache_key_to_path),
            'cache_type': 'optimized',
            'max_size': self._optimized_cache.max_size
        }


class TemplateCache:
    """
    Cache für geladene Templates um Performance zu verbessern.

    FALLBACK Implementation - wird nur genutzt wenn OptimizedTemplateCache nicht verfügbar.
    """

    def __init__(self, max_cache_size: int = 10):
        self.cache: Dict[str, Dict[str, Any]] = {}
        self.max_cache_size = max_cache_size
        self.access_times: Dict[str, datetime] = {}
        logger.debug("Using FALLBACK TemplateCache (basic implementation)")

    def get(self, cache_key: str) -> Optional[Document]:
        """Holt Template aus Cache."""
        try:
            if cache_key in self.cache:
                self.access_times[cache_key] = datetime.now()
                # Create new Document instance from cached template
                cached_template = self.cache[cache_key]['template']
                return Document(cached_template.part.package.part_related_by(cached_template.part.partname))
        except Exception as e:
            logger.warning(f"Error retrieving template from cache: {e}")
        return None

    def put(self, cache_key: str, template: Document, template_info: TemplateInfo):
        """Speichert Template in Cache."""
        try:
            # Cleanup cache if full
            if len(self.cache) >= self.max_cache_size:
                self._evict_oldest()

            # Cache template with metadata
            self.cache[cache_key] = {
                'template': template,
                'template_info': template_info,
                'cached_at': datetime.now()
            }
            self.access_times[cache_key] = datetime.now()

        except Exception as e:
            logger.warning(f"Error caching template: {e}")

    def _evict_oldest(self):
        """Entfernt ältestes Template aus Cache."""
        if not self.access_times:
            return

        oldest_key = min(self.access_times.items(), key=lambda x: x[1])[0]
        self.cache.pop(oldest_key, None)
        self.access_times.pop(oldest_key, None)

    def clear(self):
        """Leert Cache komplett."""
        self.cache.clear()
        self.access_times.clear()

    def get_stats(self) -> Dict[str, Any]:
        """Gibt Cache-Statistiken zurück."""
        return {
            'cached_templates': len(self.cache),
            'max_cache_size': self.max_cache_size,
            'cache_keys': list(self.cache.keys())
        }


class TemplateManager:
    """
    Zentrale Template-Verwaltung für DOCX-Templates.

    Ersetzt die Template-Loading-Logik aus WordTemplateService mit einem
    zentralen, cache-fähigen Template-Management-System.
    """

    def __init__(self, template_base_dir: Optional[Path] = None):
        """
        Initialize Template Manager.

        Args:
            template_base_dir: Basis-Verzeichnis für Templates (optional)
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available - Template Manager disabled")
            raise ImportError("python-docx required for template functionality")

        # Template-Verzeichnis bestimmen
        if template_base_dir is None:
            # Standard: resources/templates relativ zur Anwendung
            # Path calculation: template_manager.py -> document_generation -> services -> application -> warehouse -> src -> neu_Medealis_Archiv -> resources/templates
            self.template_base_dir = (
                Path(__file__).parent.parent.parent.parent.parent.parent / "resources" / "templates"
            )
        else:
            self.template_base_dir = template_base_dir

        # Validiere Template-Verzeichnis
        if not self.template_base_dir.exists():
            logger.error(f"Template directory not found: {self.template_base_dir}")

        # Template Registry
        self.template_registry = TEMPLATE_REGISTRY

        # Cache für Performance (optimiert wenn verfügbar)
        if USE_OPTIMIZED_CACHE:
            self.template_cache = OptimizedTemplateCacheAdapter(max_size=20)
            logger.info("TemplateManager using OPTIMIZED cache (TTL: 30min, max: 20 templates)")
        else:
            self.template_cache = TemplateCache(max_cache_size=10)
            logger.info("TemplateManager using FALLBACK cache (max: 10 templates)")

        # Statistics
        self.stats = {
            'templates_loaded': 0,
            'cache_hits': 0,
            'cache_misses': 0,
            'load_errors': 0
        }

        logger.info(f"TemplateManager initialized - Template directory: {self.template_base_dir}")
        self._validate_template_availability()

    def get_template_info(self, document_type: DocumentType) -> Optional[TemplateInfo]:
        """
        Gibt TemplateInfo für DocumentType zurück.

        Args:
            document_type: Document Type

        Returns:
            TemplateInfo oder None wenn nicht unterstützt
        """
        return get_template_info(document_type)

    def load_template(
        self,
        document_type: DocumentType,
        use_cache: bool = True
    ) -> Optional[Document]:
        """
        Lädt DOCX-Template für Document Type.

        Args:
            document_type: Document Type
            use_cache: Ob Template-Cache verwendet werden soll

        Returns:
            Geladenes Document-Objekt oder None bei Fehler

        Raises:
            FileNotFoundError: Template-Datei nicht gefunden
            Exception: Fehler beim Laden
        """
        try:
            # Template-Info holen
            template_info = self.get_template_info(document_type)
            if not template_info:
                raise ValueError(f"No template registered for document type: {document_type}")

            # Cache-Key generieren
            cache_key = f"{document_type.value}_{template_info.version}"

            # Aus Cache versuchen falls aktiviert
            if use_cache:
                cached_template = self.template_cache.get(cache_key)
                if cached_template:
                    self.stats['cache_hits'] += 1
                    logger.debug(f"Template loaded from cache: {document_type.value}")
                    return cached_template

            self.stats['cache_misses'] += 1

            # Template laden
            template = self._load_template_from_file(template_info)

            # In Cache speichern
            if use_cache and template:
                self.template_cache.put(cache_key, template, template_info)

            self.stats['templates_loaded'] += 1
            logger.info(f"Template loaded successfully: {document_type.value} from {template_info.filename}")

            return template

        except Exception as e:
            error_msg = f"Error loading template for {document_type.value}: {e}"
            logger.error(error_msg)
            self.stats['load_errors'] += 1
            raise

    def _load_template_from_file(self, template_info: TemplateInfo) -> Document:
        """
        Lädt Template aus Datei.

        Args:
            template_info: Template-Informationen

        Returns:
            Geladenes Document

        Raises:
            FileNotFoundError: Template nicht gefunden
            Exception: Fehler beim Laden
        """
        try:
            # Template-Pfad bestimmen
            template_path = self._resolve_template_path(template_info.filename)

            if not template_path.exists():
                # PyInstaller Bundle Support
                if hasattr(sys, '_MEIPASS'):
                    bundle_template_path = Path(sys._MEIPASS) / 'templates' / template_info.filename
                    if bundle_template_path.exists():
                        template_path = bundle_template_path

            if not template_path.exists():
                raise FileNotFoundError(f"Template file not found: {template_path}")

            # Template laden
            logger.debug(f"Loading template from: {template_path}")

            # Verschiedene Lade-Methoden probieren (für Kompatibilität)
            try:
                # Standardmethode
                document = Document(str(template_path))
            except Exception as load_error:
                logger.warning(f"Standard loading failed, trying alternative method: {load_error}")

                # Alternative Methode für problematische Dateien
                with open(template_path, 'rb') as template_file:
                    document = Document(template_file)

            logger.debug(f"Template loaded successfully from: {template_path}")
            return document

        except Exception as e:
            error_msg = f"Failed to load template {template_info.filename}: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg) from e

    def _resolve_template_path(self, filename: str) -> Path:
        """
        Löst Template-Pfad für Dateiname auf.

        Args:
            filename: Template-Dateiname

        Returns:
            Vollständiger Pfad zur Template-Datei
        """
        return self.template_base_dir / filename

    def validate_template(
        self,
        document: Document,
        document_type: DocumentType
    ) -> ValidationResult:
        """
        Validiert Template-Struktur und Placeholder.

        Args:
            document: Geladenes Document
            document_type: Document Type

        Returns:
            ValidationResult mit Validierungsstatus
        """
        validation_result = ValidationResult(is_valid=True)

        try:
            template_info = self.get_template_info(document_type)
            if not template_info:
                validation_result.add_error(f"No template info available for {document_type.value}")
                return validation_result

            # Template-Struktur prüfen
            if not document.paragraphs and not document.tables:
                validation_result.add_error("Template appears to be empty (no paragraphs or tables)")

            # Placeholder-Validierung
            template_text = self._extract_template_text(document)
            placeholder_validation = self._validate_template_placeholders(
                template_text, template_info
            )

            if not placeholder_validation['has_required_placeholders']:
                missing = placeholder_validation['missing_required']
                validation_result.add_warning(f"Missing required placeholders: {missing}")

            # Template-spezifische Validierung
            if template_info.has_tables and not document.tables:
                validation_result.add_warning("Template is marked as having tables but none found")

            logger.debug(f"Template validation completed for {document_type.value}")

        except Exception as e:
            validation_result.add_error(f"Template validation failed: {str(e)}")

        return validation_result

    def _extract_template_text(self, document: Document) -> str:
        """Extrahiert gesamten Text aus Template für Placeholder-Analyse."""
        try:
            all_text = []

            # Text aus Paragraphen
            for paragraph in document.paragraphs:
                all_text.append(paragraph.text)

            # Text aus Tabellen
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)

            return " ".join(all_text)

        except Exception as e:
            logger.error(f"Error extracting template text: {e}")
            return ""

    def _validate_template_placeholders(
        self,
        template_text: str,
        template_info: TemplateInfo
    ) -> Dict[str, Any]:
        """Validiert Placeholder im Template-Text."""
        try:
            # Import hier um Circular Import zu vermeiden
            from .placeholder_engine import PlaceholderEngine

            placeholder_engine = PlaceholderEngine()
            found_placeholders = placeholder_engine.find_placeholders_in_text(template_text)

            missing_required = [
                p for p in template_info.required_placeholders
                if p not in found_placeholders
            ]

            return {
                'found_placeholders': found_placeholders,
                'missing_required': missing_required,
                'has_required_placeholders': len(missing_required) == 0,
                'placeholder_count': len(found_placeholders)
            }

        except Exception as e:
            logger.error(f"Error validating placeholders: {e}")
            return {
                'found_placeholders': [],
                'missing_required': template_info.required_placeholders,
                'has_required_placeholders': False,
                'placeholder_count': 0
            }

    def get_required_placeholders(self, document_type: DocumentType) -> List[str]:
        """
        Gibt erforderliche Placeholder für Document Type zurück.

        Args:
            document_type: Document Type

        Returns:
            Liste der erforderlichen Placeholder
        """
        template_info = self.get_template_info(document_type)
        return template_info.required_placeholders if template_info else []

    def get_all_placeholders(self, document_type: DocumentType) -> List[str]:
        """
        Gibt alle Placeholder (required + optional) für Document Type zurück.

        Args:
            document_type: Document Type

        Returns:
            Liste aller Placeholder
        """
        template_info = self.get_template_info(document_type)
        return template_info.get_all_placeholders() if template_info else []

    def list_available_templates(self) -> List[Dict[str, Any]]:
        """
        Gibt Liste aller verfügbaren Templates zurück.

        Returns:
            Liste mit Template-Informationen
        """
        available_templates = []

        for document_type, template_info in self.template_registry.items():
            template_path = self._resolve_template_path(template_info.filename)

            available_templates.append({
                'document_type': document_type,
                'name': template_info.name,
                'filename': template_info.filename,
                'exists': template_path.exists(),
                'path': str(template_path),
                'required_placeholders': template_info.required_placeholders,
                'optional_placeholders': template_info.optional_placeholders,
                'supported_formats': template_info.supported_formats,
                'version': template_info.version
            })

        return available_templates

    def _validate_template_availability(self):
        """Validiert Verfügbarkeit aller registrierten Templates."""
        try:
            available_count = 0
            missing_templates = []

            for document_type, template_info in self.template_registry.items():
                template_path = self._resolve_template_path(template_info.filename)

                if template_path.exists():
                    available_count += 1
                    logger.debug(f"Template available: {template_info.filename}")
                else:
                    missing_templates.append({
                        'document_type': document_type.value,
                        'filename': template_info.filename,
                        'expected_path': str(template_path)
                    })

            logger.info(f"Template availability: {available_count}/{len(self.template_registry)} templates available")

            if missing_templates:
                logger.warning(f"Missing templates: {[t['filename'] for t in missing_templates]}")
                for template in missing_templates:
                    logger.warning(f"  - {template['filename']} expected at: {template['expected_path']}")

        except Exception as e:
            logger.error(f"Error validating template availability: {e}")

    def get_template_stats(self) -> Dict[str, Any]:
        """
        Gibt Template Manager Statistiken zurück.

        Returns:
            Dictionary mit Statistiken
        """
        cache_stats = self.template_cache.get_stats()

        return {
            'template_base_dir': str(self.template_base_dir),
            'base_dir_exists': self.template_base_dir.exists(),
            'registered_templates': len(self.template_registry),
            'loading_stats': self.stats.copy(),
            'cache_stats': cache_stats,
            'available_templates': len([
                t for t in self.list_available_templates() if t['exists']
            ])
        }

    def clear_cache(self):
        """Leert Template-Cache."""
        self.template_cache.clear()
        logger.info("Template cache cleared")

    def reload_template_registry(self):
        """
        Lädt Template-Registry neu (für Hot-Reload während Entwicklung).
        """
        try:
            # Clear cache
            self.clear_cache()

            # Reload registry from document_types module
            from .document_types import TEMPLATE_REGISTRY as NEW_REGISTRY
            self.template_registry = NEW_REGISTRY

            # Re-validate availability
            self._validate_template_availability()

            logger.info("Template registry reloaded successfully")

        except Exception as e:
            logger.error(f"Error reloading template registry: {e}")
            raise

    def create_template_backup(self, backup_dir: Path):
        """
        Erstellt Backup aller Template-Dateien.

        Args:
            backup_dir: Verzeichnis für Backup
        """
        try:
            import shutil
            from datetime import datetime

            backup_dir.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            backed_up_count = 0

            for template_info in self.template_registry.values():
                template_path = self._resolve_template_path(template_info.filename)

                if template_path.exists():
                    backup_filename = f"{timestamp}_{template_info.filename}"
                    backup_path = backup_dir / backup_filename

                    shutil.copy2(template_path, backup_path)
                    backed_up_count += 1
                    logger.debug(f"Template backed up: {template_info.filename} -> {backup_filename}")

            logger.info(f"Template backup completed: {backed_up_count} templates backed up to {backup_dir}")

        except Exception as e:
            logger.error(f"Error creating template backup: {e}")
            raise