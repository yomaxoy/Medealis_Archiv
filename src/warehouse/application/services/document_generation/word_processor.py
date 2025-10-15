"""
Word Processor - DOCX Document Processing Engine

Zentrale Engine für DOCX-Dokument-Verarbeitung:
- Template-basierte Dokument-Generierung
- Placeholder-Ersetzung in Paragraphen und Tabellen
- Dokument-Strukturierung und -Formatierung
- Integration mit PlaceholderEngine und TemplateManager
"""

import logging
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
from io import BytesIO

# External dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - Word processing disabled")

# Internal imports
from .generation_models import GenerationContext, GenerationResult, ValidationResult
from .document_types import DocumentType, get_template_info
from .placeholder_engine import PlaceholderEngine
from .template_manager import TemplateManager

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Kern-Engine für DOCX Dokument-Verarbeitung.

    Übernimmt die eigentliche Template-zu-Dokument-Transformation mit:
    - Placeholder-Ersetzung in Text und Tabellen
    - Dokument-Strukturierung
    - Formatierung und Layout
    """

    def __init__(self):
        """Initialize Document Processor."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for document processing")

        self.placeholder_engine = PlaceholderEngine()

        # Processing Statistics
        self.stats = {
            'documents_processed': 0,
            'paragraphs_processed': 0,
            'tables_processed': 0,
            'placeholders_replaced': 0,
            'processing_errors': 0
        }

        logger.info("DocumentProcessor initialized")

    def process_document(
        self,
        document: Document,
        context: GenerationContext
    ) -> Tuple[Document, Dict[str, Any]]:
        """
        Verarbeitet DOCX-Document mit GenerationContext.

        Args:
            document: Geladenes DOCX Document
            context: Generation Context mit Daten

        Returns:
            Tuple (processed_document, processing_metadata)
        """
        try:
            start_time = datetime.now()
            processing_metadata = {
                'paragraphs_processed': 0,
                'tables_processed': 0,
                'placeholders_replaced': 0,
                'processing_warnings': [],
                'processing_time': 0.0
            }

            # Context zu Placeholder Dictionary konvertieren
            placeholder_dict = context.to_placeholder_dict()

            logger.debug(f"Processing document with {len(placeholder_dict)} placeholders")

            # 1. Verarbeite Paragraphen
            paragraphs_result = self._process_paragraphs(document, placeholder_dict)
            processing_metadata['paragraphs_processed'] = paragraphs_result['count']
            processing_metadata['placeholders_replaced'] += paragraphs_result['placeholders_replaced']

            # 2. Verarbeite Tabellen
            tables_result = self._process_tables(document, placeholder_dict)
            processing_metadata['tables_processed'] = tables_result['count']
            processing_metadata['placeholders_replaced'] += tables_result['placeholders_replaced']

            # 3. Verarbeite Header/Footer (falls vorhanden)
            header_footer_result = self._process_headers_footers(document, placeholder_dict)
            processing_metadata['placeholders_replaced'] += header_footer_result['placeholders_replaced']

            # 4. Dokument-Metadaten updaten
            self._update_document_metadata(document, context)

            # Processing abschließen
            end_time = datetime.now()
            processing_metadata['processing_time'] = (end_time - start_time).total_seconds()

            # Statistics updaten
            self.stats['documents_processed'] += 1
            self.stats['paragraphs_processed'] += processing_metadata['paragraphs_processed']
            self.stats['tables_processed'] += processing_metadata['tables_processed']
            self.stats['placeholders_replaced'] += processing_metadata['placeholders_replaced']

            logger.info(f"Document processed successfully - {processing_metadata['placeholders_replaced']} placeholders replaced in {processing_metadata['processing_time']:.2f}s")

            return document, processing_metadata

        except Exception as e:
            self.stats['processing_errors'] += 1
            error_msg = f"Error processing document: {e}"
            logger.error(error_msg)
            raise Exception(error_msg) from e

    def _process_paragraphs(
        self,
        document: Document,
        placeholder_dict: Dict[str, str]
    ) -> Dict[str, Any]:
        """Verarbeitet alle Paragraphen im Dokument."""
        try:
            result = {'count': 0, 'placeholders_replaced': 0}

            for paragraph in document.paragraphs:
                if paragraph.text.strip():  # Nur nicht-leere Paragraphen
                    original_text = paragraph.text
                    processed_text = self.placeholder_engine.replace_placeholders(
                        original_text, placeholder_dict
                    )

                    if original_text != processed_text:
                        # Ersetze Text in Paragraph (behält Formatierung bei)
                        self._replace_paragraph_text(paragraph, processed_text)

                        # Zähle ersetzte Placeholders
                        result['placeholders_replaced'] += self.placeholder_engine._count_placeholders_in_text(original_text)

                    result['count'] += 1

            logger.debug(f"Processed {result['count']} paragraphs")
            return result

        except Exception as e:
            logger.error(f"Error processing paragraphs: {e}")
            return {'count': 0, 'placeholders_replaced': 0}

    def _process_tables(
        self,
        document: Document,
        placeholder_dict: Dict[str, str]
    ) -> Dict[str, Any]:
        """Verarbeitet alle Tabellen im Dokument."""
        try:
            result = {'count': 0, 'placeholders_replaced': 0}

            for table in document.tables:
                table_placeholders = 0

                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            processed_text = self.placeholder_engine.replace_placeholders(
                                original_text, placeholder_dict
                            )

                            if original_text != processed_text:
                                self._replace_paragraph_text(paragraph, processed_text)
                                table_placeholders += self.placeholder_engine._count_placeholders_in_text(original_text)

                result['placeholders_replaced'] += table_placeholders
                result['count'] += 1

                if table_placeholders > 0:
                    logger.debug(f"Processed table with {table_placeholders} placeholder replacements")

            logger.debug(f"Processed {result['count']} tables")
            return result

        except Exception as e:
            logger.error(f"Error processing tables: {e}")
            return {'count': 0, 'placeholders_replaced': 0}

    def _process_headers_footers(
        self,
        document: Document,
        placeholder_dict: Dict[str, str]
    ) -> Dict[str, Any]:
        """Verarbeitet Header und Footer."""
        try:
            result = {'placeholders_replaced': 0}

            # Header verarbeiten
            for section in document.sections:
                # Header
                if section.header:
                    for paragraph in section.header.paragraphs:
                        original_text = paragraph.text
                        processed_text = self.placeholder_engine.replace_placeholders(
                            original_text, placeholder_dict
                        )
                        if original_text != processed_text:
                            self._replace_paragraph_text(paragraph, processed_text)
                            result['placeholders_replaced'] += self.placeholder_engine._count_placeholders_in_text(original_text)

                # Footer
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        original_text = paragraph.text
                        processed_text = self.placeholder_engine.replace_placeholders(
                            original_text, placeholder_dict
                        )
                        if original_text != processed_text:
                            self._replace_paragraph_text(paragraph, processed_text)
                            result['placeholders_replaced'] += self.placeholder_engine._count_placeholders_in_text(original_text)

            if result['placeholders_replaced'] > 0:
                logger.debug(f"Processed headers/footers with {result['placeholders_replaced']} replacements")

            return result

        except Exception as e:
            logger.error(f"Error processing headers/footers: {e}")
            return {'placeholders_replaced': 0}

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """
        Ersetzt Paragraph-Text und behält Formatierung bei.

        Args:
            paragraph: python-docx Paragraph Objekt
            new_text: Neuer Text
        """
        try:
            # Behält erste Run-Formatierung bei
            if paragraph.runs:
                first_run = paragraph.runs[0]
                # Lösche alle Runs außer dem ersten
                for _ in range(len(paragraph.runs) - 1, 0, -1):
                    paragraph._element.remove(paragraph.runs[_]._element)

                # Setze neuen Text im ersten Run
                first_run.text = new_text
            else:
                # Kein Run vorhanden, erstelle neuen
                paragraph.text = new_text

        except Exception as e:
            logger.warning(f"Error replacing paragraph text: {e}")
            # Fallback: Direkter Text-Ersatz
            paragraph.text = new_text

    def _update_document_metadata(self, document: Document, context: GenerationContext):
        """Updated Dokument-Metadaten mit Context-Informationen."""
        try:
            # Core Properties setzen
            props = document.core_properties
            props.author = f"Medealis System - {context.employee_name}" if context.employee_name else "Medealis System"
            props.title = f"{context.document_type.display_name if context.document_type else 'Document'} - {context.batch_number}"
            props.subject = f"Document for Batch {context.batch_number}, Delivery {context.delivery_number}"
            props.created = context.generation_timestamp
            props.modified = context.generation_timestamp

            # Custom Properties (falls verfügbar)
            try:
                custom_props = document.custom_properties
                custom_props['batch_number'] = context.batch_number
                custom_props['delivery_number'] = context.delivery_number
                if context.article_number:
                    custom_props['article_number'] = context.article_number
                custom_props['generation_date'] = context.generation_timestamp.strftime('%Y-%m-%d %H:%M:%S')
                custom_props['template_version'] = context.template_version
            except Exception as custom_error:
                logger.debug(f"Could not set custom properties: {custom_error}")

            logger.debug("Document metadata updated")

        except Exception as e:
            logger.warning(f"Error updating document metadata: {e}")

    def get_processing_statistics(self) -> Dict[str, Any]:
        """Gibt Processing-Statistiken zurück."""
        return {
            'documents_processed': self.stats['documents_processed'],
            'paragraphs_processed': self.stats['paragraphs_processed'],
            'tables_processed': self.stats['tables_processed'],
            'placeholders_replaced': self.stats['placeholders_replaced'],
            'processing_errors': self.stats['processing_errors'],
            'placeholder_engine_stats': self.placeholder_engine.get_statistics()
        }


class WordProcessor:
    """
    High-Level Word Processing Service.

    Kombiniert TemplateManager, DocumentProcessor und file I/O für
    vollständige DOCX-Verarbeitung.
    """

    def __init__(self, template_base_dir: Optional[Path] = None):
        """
        Initialize Word Processor.

        Args:
            template_base_dir: Basis-Verzeichnis für Templates (optional)
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required for Word processing")

        # Core Components
        self.template_manager = TemplateManager(template_base_dir)
        self.document_processor = DocumentProcessor()

        # Service Statistics
        self.stats = {
            'documents_generated': 0,
            'successful_generations': 0,
            'failed_generations': 0,
            'total_processing_time': 0.0
        }

        logger.info("WordProcessor initialized")

    def generate_document_bytes(
        self,
        context: GenerationContext
    ) -> Tuple[Optional[bytes], Dict[str, Any]]:
        """
        Generiert DOCX-Dokument als bytes (ohne lokale Speicherung).

        Args:
            context: Generation Context mit allen Daten

        Returns:
            Tuple (document_bytes, metadata_dict)
            - document_bytes: DOCX als bytes oder None bei Fehler
            - metadata_dict: Metadaten zur Generierung
        """
        metadata = {
            'success': False,
            'error': None,
            'template_used': None,
            'template_version': None,
            'placeholders_replaced': 0,
            'context_completeness': context.completeness_score,
            'generation_time': 0.0
        }

        try:
            start_time = datetime.now()

            # Validiere Context
            if not context.document_type:
                raise ValueError("document_type is required in GenerationContext")

            # 1. Template laden
            logger.debug(f"Loading template for {context.document_type.value}")
            template_document = self.template_manager.load_template(context.document_type)

            if not template_document:
                raise ValueError(f"Could not load template for {context.document_type.value}")

            template_info = self.template_manager.get_template_info(context.document_type)
            metadata['template_used'] = template_info.filename if template_info else "unknown"
            metadata['template_version'] = template_info.version if template_info else "unknown"

            # 2. Dokument verarbeiten
            logger.debug(f"Processing document with context completeness: {context.completeness_score:.1%}")
            processed_document, processing_metadata = self.document_processor.process_document(
                template_document, context
            )

            metadata['placeholders_replaced'] = processing_metadata['placeholders_replaced']

            # 3. Dokument als bytes speichern
            doc_bytes_io = BytesIO()
            processed_document.save(doc_bytes_io)
            doc_bytes = doc_bytes_io.getvalue()

            # Erfolg
            metadata['success'] = True
            end_time = datetime.now()
            metadata['generation_time'] = (end_time - start_time).total_seconds()

            logger.debug(f"Document generated as bytes: {len(doc_bytes)} bytes, {metadata['placeholders_replaced']} placeholders")

            return doc_bytes, metadata

        except Exception as e:
            error_msg = f"Failed to generate document bytes for {context.document_type.value if context.document_type else 'unknown'}: {str(e)}"
            logger.error(error_msg)
            metadata['error'] = error_msg
            return None, metadata

    def generate_document(
        self,
        context: GenerationContext,
        output_path: Optional[Path] = None
    ) -> GenerationResult:
        """
        Generiert DOCX-Dokument aus GenerationContext.

        DEPRECATED: Diese Methode speichert direkt lokal.
        Verwende stattdessen generate_document_bytes() + DocumentStorageService.

        Args:
            context: Generation Context mit allen Daten
            output_path: Ausgabepfad (optional)

        Returns:
            GenerationResult mit Erfolg/Fehler-Informationen
        """
        start_time = datetime.now()
        generation_result = GenerationResult(success=False)

        try:
            # Validiere Context
            if not context.document_type:
                raise ValueError("document_type is required in GenerationContext")

            generation_result.document_type = context.document_type

            # 1. Template laden
            logger.info(f"Loading template for {context.document_type.value}")
            template_document = self.template_manager.load_template(context.document_type)

            if not template_document:
                raise ValueError(f"Could not load template for {context.document_type.value}")

            template_info = self.template_manager.get_template_info(context.document_type)
            generation_result.template_used = template_info.filename if template_info else "unknown"
            generation_result.template_version = template_info.version if template_info else "unknown"

            # 2. Template validieren
            template_validation = self.template_manager.validate_template(template_document, context.document_type)
            if not template_validation.is_valid:
                for error in template_validation.errors:
                    generation_result.add_warning(f"Template validation: {error}")

            # 3. Dokument verarbeiten
            logger.info(f"Processing document with context completeness: {context.completeness_score:.1%}")
            processed_document, processing_metadata = self.document_processor.process_document(
                template_document, context
            )

            # 4. Ausgabepfad bestimmen
            if not output_path:
                output_path = self._generate_output_path(context)

            # Stelle sicher dass Ausgabeverzeichnis existiert
            output_path.parent.mkdir(parents=True, exist_ok=True)

            # 5. Dokument speichern
            logger.info(f"Saving document to {output_path}")
            processed_document.save(str(output_path))

            # 6. Erfolg setzen
            generation_result.success = True
            generation_result.document_path = output_path
            generation_result.placeholder_count = processing_metadata['placeholders_replaced']
            generation_result.context_completeness = context.completeness_score

            # Metadaten hinzufügen
            generation_result.metadata.update({
                'processing_metadata': processing_metadata,
                'template_info': template_info.__dict__ if template_info else {},
                'context_source': context.context_source,
                'generation_timestamp': context.generation_timestamp.isoformat()
            })

            # Statistics
            self.stats['documents_generated'] += 1
            self.stats['successful_generations'] += 1

            end_time = datetime.now()
            generation_result.generation_time = (end_time - start_time).total_seconds()
            self.stats['total_processing_time'] += generation_result.generation_time

            logger.info(f"Document generated successfully: {output_path}")

            return generation_result

        except Exception as e:
            # Fehler-Handling
            error_msg = f"Failed to generate document for {context.document_type.value if context.document_type else 'unknown'}: {str(e)}"
            logger.error(error_msg)

            generation_result.set_error(error_msg)

            end_time = datetime.now()
            generation_result.generation_time = (end_time - start_time).total_seconds()

            # Statistics
            self.stats['documents_generated'] += 1
            self.stats['failed_generations'] += 1
            self.stats['total_processing_time'] += generation_result.generation_time

            return generation_result

    def generate_batch_documents(
        self,
        contexts: List[GenerationContext],
        output_base_path: Optional[Path] = None
    ) -> List[GenerationResult]:
        """
        Generiert mehrere Dokumente in einem Batch.

        Args:
            contexts: Liste von GenerationContexts
            output_base_path: Basis-Pfad für Ausgabe (optional)

        Returns:
            Liste von GenerationResults
        """
        results = []

        try:
            logger.info(f"Starting batch generation of {len(contexts)} documents")

            for i, context in enumerate(contexts, 1):
                logger.info(f"Processing document {i}/{len(contexts)}: {context.document_type.value if context.document_type else 'unknown'}")

                # Output path für dieses Dokument
                document_output_path = None
                if output_base_path:
                    filename = self._generate_filename(context)
                    document_output_path = output_base_path / filename

                # Einzelnes Dokument generieren
                result = self.generate_document(context, document_output_path)
                results.append(result)

                # Log Zwischenergebnis
                if result.success:
                    logger.debug(f"Document {i} generated successfully")
                else:
                    logger.warning(f"Document {i} generation failed: {result.error}")

            successful_count = sum(1 for r in results if r.success)
            logger.info(f"Batch generation completed: {successful_count}/{len(results)} documents successful")

            return results

        except Exception as e:
            error_msg = f"Batch generation failed: {e}"
            logger.error(error_msg)

            # Return error result for any remaining contexts
            for context in contexts[len(results):]:
                error_result = GenerationResult(success=False)
                error_result.set_error(error_msg)
                error_result.document_type = context.document_type
                results.append(error_result)

            return results

    def _generate_output_path(self, context: GenerationContext) -> Path:
        """Generiert Standard-Ausgabepfad für Context."""
        try:
            # FIXED: Versuche zuerst echten Artikelordner zu verwenden
            if (context.batch_number and context.delivery_number and
                hasattr(context, 'article_number') and context.article_number):
                try:
                    # Import DocumentStorageService für echte Pfad-Auflösung
                    from ..document_storage.document_storage_service import document_storage_service

                    # Hole korrekten Artikelordner-Pfad
                    storage_path, warnings = document_storage_service.get_document_path(
                        batch_number=context.batch_number,
                        delivery_number=context.delivery_number,
                        article_number=getattr(context, 'article_number', ''),
                        supplier_name=getattr(context, 'supplier_name', ''),
                        create_folders=True
                    )

                    if storage_path and storage_path != Path("."):
                        # Dateiname generieren
                        filename = self._generate_filename(context)
                        return storage_path / filename

                except Exception as storage_error:
                    logger.warning(f"Could not resolve storage path, using temp: {storage_error}")

            # Fallback: Basis-Verzeichnis (temp oder documents)
            base_dir = Path.cwd() / "temp" / "generated_documents"

            # Dateiname generieren
            filename = self._generate_filename(context)

            return base_dir / filename

        except Exception as e:
            logger.warning(f"Error generating output path: {e}")
            # Fallback
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return Path.cwd() / "temp" / f"document_{timestamp}.docx"

    def _generate_filename(self, context: GenerationContext) -> str:
        """Generiert Dateiname für Context."""
        try:
            if context.document_type:
                template_info = self.template_manager.get_template_info(context.document_type)
                if template_info:
                    return template_info.generate_filename(context.to_placeholder_dict())

            # Fallback
            timestamp = context.generation_timestamp.strftime("%Y%m%d_%H%M%S")
            doc_type = context.document_type.value if context.document_type else "document"
            return f"{doc_type}_{context.batch_number}_{timestamp}.docx"

        except Exception as e:
            logger.warning(f"Error generating filename: {e}")
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            return f"document_{timestamp}.docx"

    def validate_processing_capabilities(self) -> Dict[str, Any]:
        """Validiert Word Processor Funktionen."""
        try:
            validation_result = {
                'docx_available': DOCX_AVAILABLE,
                'template_manager_ready': False,
                'document_processor_ready': False,
                'available_templates': 0,
                'processing_stats': self.stats,
                'errors': []
            }

            # Template Manager prüfen
            try:
                template_stats = self.template_manager.get_template_stats()
                validation_result['template_manager_ready'] = template_stats['base_dir_exists']
                validation_result['available_templates'] = template_stats['available_templates']
            except Exception as e:
                validation_result['errors'].append(f"Template Manager error: {e}")

            # Document Processor prüfen
            try:
                processor_stats = self.document_processor.get_processing_statistics()
                validation_result['document_processor_ready'] = True
            except Exception as e:
                validation_result['errors'].append(f"Document Processor error: {e}")

            return validation_result

        except Exception as e:
            return {
                'docx_available': False,
                'error': str(e)
            }

    def get_word_processor_statistics(self) -> Dict[str, Any]:
        """Gibt vollständige Word Processor Statistiken zurück."""
        return {
            'word_processor_stats': self.stats,
            'template_manager_stats': self.template_manager.get_template_stats(),
            'document_processor_stats': self.document_processor.get_processing_statistics()
        }