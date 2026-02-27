"""
Hilfsskript zum Aktualisieren des Wareneingangskontrolle-Templates
Ersetzt [[itemnr]] durch [[artikel]] im Word-Dokument
"""
from docx import Document
from pathlib import Path

template_path = Path(r"C:\Users\krueg\OneDrive\Desktop\Medealis Archiv\resources\templates\Fo0113_Wareneingangskontrolle.docx")
backup_path = template_path.with_suffix('.docx.backup')

print(f"Lade Template: {template_path}")

# Backup erstellen
import shutil
shutil.copy2(template_path, backup_path)
print(f"Backup erstellt: {backup_path}")

# Template laden
doc = Document(template_path)

# Durchsuche alle Tabellen
replacements_made = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if '[[itemnr]]' in cell.text:
                print(f"Gefunden in Zelle: '{cell.text}'")
                # Ersetze in allen Paragraphs der Zelle
                for paragraph in cell.paragraphs:
                    if '[[itemnr]]' in paragraph.text:
                        # Durchsuche alle Runs
                        for run in paragraph.runs:
                            if '[[itemnr]]' in run.text:
                                run.text = run.text.replace('[[itemnr]]', '[[artikel]]')
                                replacements_made += 1
                                print(f"Ersetzt: [[itemnr]] -> [[artikel]]")

# Speichern
if replacements_made > 0:
    doc.save(template_path)
    print(f"\nTemplate aktualisiert! ({replacements_made} Ersetzungen)")
    print(f"Gespeichert: {template_path}")
else:
    print("\nKeine [[itemnr]] Platzhalter gefunden!")
